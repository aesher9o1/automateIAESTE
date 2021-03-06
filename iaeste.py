# -*- coding: utf-8 -*-
"""
Created on Mon Feb  4 15:51:12 2019

@author: aesher9o1
"""

import pandas as pd
from docx import *
import os

filename = ""

array_files=os.listdir()
for i in array_files:
    if(".xlsx" in i):
        filename= i.replace(".xlsx",'')
        
        if not os.path.exists("factsheet"):
            os.makedirs("factsheet")

        xl_file = pd.read_excel(filename+".xlsx")
        xl_file= xl_file.drop(['Timestamp'], axis=1)
        xl_file= xl_file.drop(['Passport'], axis=1)
        
        
        
        "Prepare document for evaluation sheet"
        document = Document()
        table = document.add_table(rows=1,cols=7 )
        table.style = 'TableGrid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'S.No'
        hdr_cells[1].text = 'IAESTE ID'
        hdr_cells[2].text = 'CGPA (out of 20)'
        hdr_cells[3].text = 'Technical Skills (out of 30)'
        hdr_cells[4].text = 'Extra Cirricular (out of 20)'
        hdr_cells[5].text = 'Aptness (out of 20)'
        hdr_cells[6].text = 'Total'
        document.add_page_break()
      
        

        """Loop to generate factsheet"""
        for index, row in xl_file.iterrows():
            p = document.add_paragraph()
            p.add_run("IAESTE ID:   ").bold = True
            p.add_run(str(row[0]))
            
            p = document.add_paragraph()
            p.add_run("Branch:   ").bold = True
            p.add_run(str(row[1]))
            
            p = document.add_paragraph()
            p.add_run("Number of Backlogs:  ").bold = True         
            p.add_run(str(row[2]))

            p = document.add_paragraph()
            p.add_run("Year of Study:   ").bold = True
            p.add_run(str(row[3]))
            
            p = document.add_paragraph()
            p.add_run("CGPA:  ").bold = True
            p.add_run(str(row[4]))
            
            document.add_heading('\nTechnical Skills:', level=0)
            document.add_paragraph(str(row[5]).strip())
            
            document.add_heading('\nExtra-Curricular Activities & Skills:', level=0)
            document.add_paragraph(str(row[6]).strip())
            
            document.add_heading('\nAptness:', level=0)
            document.add_paragraph(str(row[7]).strip())
            document.add_page_break()
            
            row_cells = table.add_row().cells
            row_cells[0].text = str(index+1)
            row_cells[1].text = str(row[0])
            
            

        document.save('./factsheet/'+filename+'.docx')
